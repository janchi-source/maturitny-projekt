from pathlib import Path
import hashlib
from uuid import uuid4

from docx import Document as DocxDocument
from PyPDF2 import PdfReader
from werkzeug.utils import secure_filename

from ..extensions import db
from ..models.document import Document, DocumentRevision, DocumentType


def extract_text_from_pdf(filepath):
    reader = PdfReader(filepath)
    chunks = []
    for page in reader.pages:
        chunks.append(page.extract_text() or "")
    return "\n".join(chunks).strip()


def extract_text_from_docx(filepath):
    doc = DocxDocument(filepath)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs).strip()


def process_upload(file, project_id, user_id, tags=None, upload_root="uploads"):
    original_name = secure_filename(file.filename or "")
    if not original_name:
        raise ValueError("Invalid filename")

    extension = Path(original_name).suffix.lower()
    if extension not in {".pdf", ".docx"}:
        raise ValueError("Unsupported file type")

    file_type = DocumentType.PDF if extension == ".pdf" else DocumentType.DOCX

    safe_stem = Path(original_name).stem or "document"
    stored_name = f"{safe_stem}_{uuid4().hex[:8]}{extension}"
    upload_dir = Path(upload_root)
    upload_dir.mkdir(parents=True, exist_ok=True)
    file_path = upload_dir / stored_name
    file.save(file_path)

    extracted_text = ""
    if file_type == DocumentType.PDF:
        extracted_text = extract_text_from_pdf(str(file_path))
    elif file_type == DocumentType.DOCX:
        extracted_text = extract_text_from_docx(str(file_path))

    normalized_tags = [tag.strip().lower() for tag in (tags or []) if tag and tag.strip()]
    version = (
        Document.query.filter_by(project_id=project_id, original_name=original_name).count() + 1
    )
    file_hash = _sha256_for_file(file_path)

    document = Document(
        filename=stored_name,
        original_name=original_name,
        file_type=file_type,
        file_size=file_path.stat().st_size,
        extracted_text=extracted_text,
        tags=normalized_tags,
        project_id=project_id,
        uploaded_by=user_id,
        updated_by=user_id,
        version=version,
    )
    db.session.add(document)
    db.session.flush()

    revision = DocumentRevision(
        document_id=document.id,
        version=version,
        filename=stored_name,
        file_size=document.file_size,
        file_hash=file_hash,
        created_by=user_id,
        change_note="Initial upload" if version == 1 else "New upload revision",
    )
    db.session.add(revision)
    db.session.commit()
    return document


def _sha256_for_file(file_path):
    hasher = hashlib.sha256()
    with file_path.open("rb") as file_handle:
        while True:
            chunk = file_handle.read(8192)
            if not chunk:
                break
            hasher.update(chunk)
    return hasher.hexdigest()
